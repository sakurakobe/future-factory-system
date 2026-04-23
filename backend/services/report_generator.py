"""
================================================================================
诊断评估报告生成器 (DOCX)
================================================================================
功能：
    根据项目的评估数据，自动生成符合标准格式的 Word 诊断报告。

报告结构（按示例模板生成）：
    1. 标题：{企业名称}未来工厂建设诊断评估报告
    2. 整体情况（概述）：
       - 评估背景说明
       - 现场调研说明
       - 总体评分（总分、满分、得分率）
       - 得分汇总表（方面 + 子类 + 满分 + 得分）
    3. 要求细项（详细分析）：
       - 按方面 → 子类 → 细项展开
       - 每子类包含：要求解读、企业现状与差距、详细评测表格
       - 评测表格列：能力域、评测题目、企业现状与差距、对应分值、
         对标未来工厂建议的提升点、目标分值
    4. 改造投入建议：
       - 按类别分组（必要投入/建议投入/待优化/已明确）
       - 投资项目汇总表

使用方式：
    generator = ReportGenerator(project, db)
    generator.generate("output.docx")

依赖：
    python-docx 库，用于 Word 文档的创建和格式化
================================================================================
"""
import json
import os
from datetime import datetime
from sqlalchemy.orm import Session
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from models.category import MajorCategory, SubCategory, SubSubCategory
from models.question import Question
from models.answer import AssessmentAnswer
from models.note import MajorCategoryNote
from models.investment import InvestmentItem
from models.supplementary import SupplementaryInfo
from config import UPLOAD_DIR
from services.report_templates import (
    SUBCATEGORY_REQUIREMENTS,
    INTELLIGENT_PRODUCTION_SUB_REQUIREMENTS,
    MAJOR_CATEGORY_DESCRIPTIONS,
    OVERVIEW_BACKGROUND,
    REQUIREMENT_INTRO,
)

LEVEL_ORDER = ['A', 'B', 'C', 'D', 'E', 'F']


class ReportGenerator:

    def __init__(self, project, db: Session):
        self.project = project
        self.db = db
        self.doc = Document()
        self._setup_styles()
        self.company_type = project.company_type or "通用"

    def _should_show_question(self, question) -> bool:
        if self.company_type == "通用":
            return True
        title = question.title
        if self.company_type == "离散":
            return "流程行业：" not in title
        if self.company_type == "流程":
            return "离散行业：" not in title
        return True

    def _setup_styles(self):
        style = self.doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(12)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    def _set_cell_font(self, cell, text, bold=False, size=10):
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.name = '宋体'
        run.font.bold = bold
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    def generate(self, filepath: str, use_ai: bool = False):
        self._use_ai = use_ai
        self._add_title()
        self._add_overview()
        self._add_detailed_analysis()
        self._add_investment_plan()
        self.doc.save(filepath)
        return filepath

    def generate_stream(self, filepath: str, use_ai: bool, callback):
        """
        生成报告并通过 callback 流式推送进度。
        callback(step, pct, ai) -> None
        """
        self._use_ai = use_ai
        callback("开始生成报告...", 0, False)
        callback("正在计算得分汇总...", 3, False)

        self._add_title()
        callback("标题已生成", 5, False)

        self._add_overview()
        callback("整体情况概览完成", 10, False)

        callback("正在生成详细分析...", 12, False)
        self._add_detailed_analysis_stream(callback)

        self._add_investment_plan()
        callback("改造投入建议完成", 95, False)

        self.doc.save(filepath)
        callback("报告文档已保存...", 100, False)

    # ------------------------------------------------------------------
    # 标题
    # ------------------------------------------------------------------
    def _add_title(self):
        self.doc.add_paragraph(
            f'{self.project.company_name}未来工厂建设诊断评估报告',
            style='Heading 1'
        ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ------------------------------------------------------------------
    # 整体情况
    # ------------------------------------------------------------------
    def _add_overview(self):
        self.doc.add_heading('整体情况', level=1)

        # ---- 背景说明 ----
        company_name_short = self.project.company_name
        if '有限公司' in company_name_short:
            company_name_short = company_name_short.split('有限公司')[0]
        if '公司' in company_name_short:
            company_name_short = company_name_short.split('公司')[0]

        now = datetime.now()
        background = OVERVIEW_BACKGROUND.format(
            company_name=self.project.company_name,
            company_name_short=company_name_short,
            year=now.year,
            month=now.month,
        )
        self.doc.add_paragraph(background)

        # ---- 调研说明 ----
        assessors = self.project.assessors or '评估专家组'
        start = self.project.assessment_start_date or '——'
        end = self.project.assessment_end_date or '——'
        self.doc.add_paragraph(
            f'现场调研期间（{start}至{end}），由{assessors}组成的诊断评估组'
            f'围绕企业总体架构、基础支撑、应用场景等主要方面，开展全面的诊断评估工作，'
            f'系统梳理企业未来工厂发展的实际水平、发展规划，'
            f'诊断未来工厂建设的问题，对未来工厂建设提出专业指导建议。'
        )

        # ---- 总体评分 ----
        answers = self.db.query(AssessmentAnswer).filter(
            AssessmentAnswer.project_id == self.project.id
        ).all()

        total_score = 0
        total_max = 0
        for a in answers:
            question = self.db.query(Question).filter(Question.id == a.question_id).first()
            if not question or not self._should_show_question(question):
                continue
            if a.score:
                total_score += a.score
            total_max += question.max_score

        percentage = (total_score / total_max * 100) if total_max > 0 else 0

        # 判断整体水平阶段
        if percentage >= 80:
            stage = "处于未来工厂水平"
        elif percentage >= 60:
            stage = "处于智能工厂向未来工厂过渡阶段"
        elif percentage >= 40:
            stage = "处于省级智能工厂培育阶段"
        else:
            stage = "处于数字化车间培育阶段"

        self.doc.add_paragraph(
            f'综合调研情况，{self.project.company_name}初步得分为{total_score:.2f}分'
            f'（满分{total_max}分），得分率{percentage:.2f}%，整体建设水平{stage}。'
        )

        # ---- 得分汇总表 ----
        majors = self.db.query(MajorCategory).order_by(MajorCategory.sort_order).all()
        subs = self.db.query(SubCategory).order_by(SubCategory.sort_order).all()

        answer_map = {a.question_id: a for a in answers}

        table = self.doc.add_table(rows=1, cols=5, style='Light Grid Accent 1')
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ['能力要素', '能力域', '总分', '得分', '得分率']
        for i, h in enumerate(headers):
            self._set_cell_font(table.rows[0].cells[i], h, bold=True)

        for major in majors:
            major_subs = [s for s in subs if s.major_category_id == major.id]
            first = True
            for sub in major_subs:
                sub_score = 0
                sub_max = 0
                subsubs = self.db.query(SubSubCategory).filter(
                    SubSubCategory.sub_category_id == sub.id
                ).all()
                for ss in subsubs:
                    questions = self.db.query(Question).filter(
                        Question.sub_sub_category_id == ss.id
                    ).all()
                    for q in questions:
                        if not self._should_show_question(q):
                            continue
                        sub_max += q.max_score
                        if q.id in answer_map and answer_map[q.id].score:
                            sub_score += answer_map[q.id].score

                row = table.add_row()
                sub_pct = (sub_score / sub_max * 100) if sub_max > 0 else 0
                if first:
                    self._set_cell_font(row.cells[0], major.name, bold=True)
                    first = False
                self._set_cell_font(row.cells[1], sub.name)
                self._set_cell_font(row.cells[2], f'{sub_max}')
                self._set_cell_font(row.cells[3], f'{sub_score:.2f}')
                self._set_cell_font(row.cells[4], f'{sub_pct:.2f}%')

        self.doc.add_paragraph('')

    # ------------------------------------------------------------------
    # 要求细项
    # ------------------------------------------------------------------
    def _add_detailed_analysis(self):
        self.doc.add_heading('要求细项', level=1)
        self.doc.add_paragraph(REQUIREMENT_INTRO)
        self._add_detailed_sections()

    def _add_detailed_analysis_stream(self, callback):
        """带进度回调的详细分析生成"""
        self.doc.add_heading('要求细项', level=1)
        self.doc.add_paragraph(REQUIREMENT_INTRO)
        self._add_detailed_sections(callback=callback)

    def _add_detailed_sections(self, callback=None):
        """详细分析的核心逻辑，支持可选进度回调"""
        majors = self.db.query(MajorCategory).order_by(MajorCategory.sort_order).all()
        answers = self.db.query(AssessmentAnswer).filter(
            AssessmentAnswer.project_id == self.project.id
        ).all()
        answer_map = {a.question_id: a for a in answers}

        total_majors = len(majors)
        # Calculate total subcategories for progress tracking
        total_subs = 0
        for major in majors:
            subs = self.db.query(SubCategory).filter(
                SubCategory.major_category_id == major.id
            ).all()
            total_subs += len(subs)

        sub_idx = 0
        pct_base = 12  # Start progress at 12%
        pct_range = 80  # Use 80% of progress (12-92%)

        for mi, major in enumerate(majors):
            if callback:
                callback(f"正在生成 [{major.name}] 概述...", pct_base + int(mi / total_majors * pct_range * 0.15), False)

            self.doc.add_heading(major.name, level=2)

            # Major Category overview paragraph
            major_score = 0
            major_max = 0
            subs = self.db.query(SubCategory).filter(
                SubCategory.major_category_id == major.id
            ).order_by(SubCategory.sort_order).all()
            sub_names = []
            for sub in subs:
                sub_score = 0
                sub_max = 0
                subsubs = self.db.query(SubSubCategory).filter(
                    SubSubCategory.sub_category_id == sub.id
                ).all()
                for ss in subsubs:
                    questions = self.db.query(Question).filter(
                        Question.sub_sub_category_id == ss.id
                    ).all()
                    for q in questions:
                        if not self._should_show_question(q):
                            continue
                        sub_max += q.max_score
                        major_max += q.max_score
                        if q.id in answer_map and answer_map[q.id].score:
                            sub_score += answer_map[q.id].score
                            major_score += answer_map[q.id].score
                sub_names.append(sub.name)

            major_pct = (major_score / major_max * 100) if major_max > 0 else 0
            major_desc = MAJOR_CATEGORY_DESCRIPTIONS.get(major.name, '')
            sub_names_str = '、'.join(sub_names)
            overview_text = (
                f'{major.name}能力要素总分{major_max}分，企业得分{major_score:.2f}分，'
                f'得分率{major_pct:.2f}%。{major_desc}涵盖{sub_names_str}等能力域。'
            )
            self.doc.add_paragraph(overview_text)

            for sub in subs:
                sub_idx += 1
                if callback:
                    pct = pct_base + int(sub_idx / total_subs * pct_range)
                    callback(f"正在分析 [{sub.name}] — 企业现状与差距...", pct, self._use_ai)

                sub_score = 0
                sub_max = 0
                target_score = 0
                subsubs = self.db.query(SubSubCategory).filter(
                    SubSubCategory.sub_category_id == sub.id
                ).all()
                for ss in subsubs:
                    questions = self.db.query(Question).filter(
                        Question.sub_sub_category_id == ss.id
                    ).all()
                    for q in questions:
                        if not self._should_show_question(q):
                            continue
                        sub_max += q.max_score
                        if q.id in answer_map:
                            a = answer_map[q.id]
                            if a.score:
                                sub_score += a.score
                            if a.target_level:
                                options = json.loads(q.options_json)
                                for opt in options:
                                    if opt['level'] == a.target_level:
                                        target_score += opt['score']
                                        break

                pct = (sub_score / sub_max * 100) if sub_max > 0 else 0

                # 子类标题
                self.doc.add_heading(sub.name, level=3)
                self.doc.add_paragraph(
                    f'{sub.name}（总分{sub_max}分，当前得分{sub_score:.2f}分，'
                    f'得分率{pct:.2f}%，目标得分{target_score:.2f}分）'
                )

                # ---- 要求解读 ----
                requirement = self._get_requirement_interpretation(sub)
                self.doc.add_paragraph(f'要求解读：{requirement}')

                # ---- 企业现状与差距 ----
                status_text = self._generate_status_text(sub, answer_map)
                self.doc.add_paragraph(f'企业现状与差距分析：{status_text}')

                # ---- 评测表格 ----
                self._add_detail_table(sub, answer_map, callback)

        if callback:
            callback("详细分析生成完毕", pct_base + pct_range, False)

    def _get_requirement_interpretation(self, sub_category) -> str:
        """获取子类的要求解读文本"""
        # 优先使用模板
        text = SUBCATEGORY_REQUIREMENTS.get(sub_category.name, '')
        if text:
            return text
        # 兜底使用数据库中的描述
        if sub_category.description:
            return sub_category.description
        return '暂无要求解读。'

    def _generate_status_text(self, sub_category, answer_map) -> str:
        """
        结构化呈现子类下各 SubSubCategory 的企业现状。
        按 SubSubCategory 分组，每个子域列出关键信息。
        如果 use_ai 开启，则尝试调用 AI 生成叙述性分析。
        """
        subsubs = self.db.query(SubSubCategory).filter(
            SubSubCategory.sub_category_id == sub_category.id
        ).order_by(SubSubCategory.sort_order).all()

        # 收集所有题目的现状数据
        status_data = []
        for ss in subsubs:
            questions = self.db.query(Question).filter(
                Question.sub_sub_category_id == ss.id
            ).order_by(Question.sort_order).all()
            for q in questions:
                if not self._should_show_question(q):
                    continue
                if q.id in answer_map:
                    a = answer_map[q.id]
                    text = a.company_status or a.communication_content
                    status_data.append({
                        "sub_sub_name": ss.name,
                        "question_title": q.title,
                        "company_status": text or "",
                        "score": a.score or 0,
                        "selected_level": a.selected_level or "",
                        "target_level": a.target_level or "",
                    })

        # 尝试 AI 增强
        if self._use_ai and status_data:
            try:
                from services.ai_enhancer import generate_status_analysis
                requirement = self._get_requirement_interpretation(sub_category)
                ai_result = generate_status_analysis(
                    requirement, status_data, self.project.company_name
                )
                if ai_result:
                    return ai_result
            except Exception:
                pass  # 降级到结构化文本

        # 结构化呈现
        parts = []
        for ss in subsubs:
            questions = self.db.query(Question).filter(
                Question.sub_sub_category_id == ss.id
            ).order_by(Question.sort_order).all()
            status_items = []
            for q in questions:
                if not self._should_show_question(q):
                    continue
                if q.id in answer_map:
                    a = answer_map[q.id]
                    text = a.company_status or a.communication_content
                    if text:
                        status_items.append(f'{q.title}：{text}')
            if status_items:
                joined = '；'.join(status_items)
                parts.append(f'{ss.name}方面，{joined}')
            else:
                parts.append(f'{ss.name}方面，暂无记录。')

        if parts:
            return '；'.join(parts) + '。'
        return '暂无记录。'

    # ------------------------------------------------------------------
    # 评测表格
    # ------------------------------------------------------------------
    def _add_detail_table(self, sub_category, answer_map, callback=None):
        subsubs = self.db.query(SubSubCategory).filter(
            SubSubCategory.sub_category_id == sub_category.id
        ).order_by(SubSubCategory.sort_order).all()

        table = self.doc.add_table(rows=1, cols=6, style='Light Grid Accent 1')
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ['能力域', '评测题目', '企业现状与差距', '对应分值',
                   '对标未来工厂建议的提升点', '目标分值']
        for i, h in enumerate(headers):
            self._set_cell_font(table.rows[0].cells[i], h, bold=True, size=9)

        q_count = 0
        for ss in subsubs:
            questions = self.db.query(Question).filter(
                Question.sub_sub_category_id == ss.id
            ).order_by(Question.sort_order).all()

            for q in questions:
                if not self._should_show_question(q):
                    continue
                q_count += 1
                row = table.add_row()
                self._set_cell_font(row.cells[0], ss.name, size=9)
                self._set_cell_font(row.cells[1], q.title, size=9)

                if q.id in answer_map:
                    a = answer_map[q.id]
                    self._set_cell_font(row.cells[2],
                                        a.company_status or a.communication_content or '', size=9)
                    self._set_cell_font(row.cells[3], f'{a.score or 0:.2f}', size=9)

                    suggestion = self._generate_suggestion(q, a)
                    self._set_cell_font(row.cells[4], suggestion, size=9)

                    if a.target_level:
                        options = json.loads(q.options_json)
                        for opt in options:
                            if opt['level'] == a.target_level:
                                self._set_cell_font(row.cells[5], f'{opt["score"]:.2f}', size=9)
                                break
                    else:
                        self._set_cell_font(row.cells[5], '', size=9)
                else:
                    self._set_cell_font(row.cells[2], '', size=9)
                    self._set_cell_font(row.cells[3], '0.00', size=9)
                    self._set_cell_font(row.cells[4], '', size=9)
                    self._set_cell_font(row.cells[5], '', size=9)

    def _generate_suggestion(self, question, answer) -> str:
        """
        根据当前等级和目标等级生成改进建议。
        如果 use_ai 开启，尝试调用 AI 生成更专业的建议。
        """
        if not answer.selected_level or not answer.target_level:
            return '——'
        if answer.selected_level == answer.target_level:
            return '已达目标，持续优化。'

        # 尝试 AI 增强
        if self._use_ai:
            try:
                from services.ai_enhancer import generate_suggestion
                current_idx = LEVEL_ORDER.index(answer.selected_level) if answer.selected_level in LEVEL_ORDER else 0
                target_idx = LEVEL_ORDER.index(answer.target_level) if answer.target_level in LEVEL_ORDER else 0

                options = json.loads(question.options_json)
                level_descs = []
                for i in range(current_idx + 1, target_idx + 1):
                    for opt in options:
                        if opt['level'] == LEVEL_ORDER[i]:
                            level_descs.append({
                                "level": opt['level'],
                                "score": opt.get('score', 0),
                                "description": opt.get('description', ''),
                            })
                            break

                if level_descs:
                    ai_result = generate_suggestion(
                        question_title=question.title,
                        company_status=answer.company_status or answer.communication_content or '',
                        current_level=answer.selected_level,
                        target_level=answer.target_level,
                        level_descriptions=level_descs,
                    )
                    if ai_result:
                        return ai_result
            except Exception:
                pass

        # 降级：提取等级描述的完整段落
        current_idx = LEVEL_ORDER.index(answer.selected_level) if answer.selected_level in LEVEL_ORDER else 0
        target_idx = LEVEL_ORDER.index(answer.target_level) if answer.target_level in LEVEL_ORDER else 0

        options = json.loads(question.options_json)
        suggestions = []
        for i in range(current_idx + 1, target_idx + 1):
            for opt in options:
                if opt['level'] == LEVEL_ORDER[i] and opt.get('description'):
                    desc = opt['description']
                    if len(desc) > 200:
                        desc = desc[:200]
                    level_label = opt.get('level', '')
                    score = opt.get('score', 0)
                    suggestions.append(f'{level_label}级({score}分)：{desc}')
                    break

        return '；'.join(suggestions) if suggestions else '——'

    # ------------------------------------------------------------------
    # 改造投入建议
    # ------------------------------------------------------------------
    def _add_investment_plan(self):
        self.doc.add_heading('改造投入建议', level=1)

        categories = {
            'necessary': '必要投入/改造项目',
            'recommended': '建议投入项目',
            'system_upgrade': '系统已有待优化项目',
            'confirmed': '已明确投入项目',
        }

        # 各类别说明
        category_intro = {
            'necessary': '此类项目为未来工厂认定关键支撑项，直接影响核心能力域得分，需优先推进落地。',
            'recommended': '此类项目为未来工厂建设重要支撑项，建议结合企业发展战略，分阶段逐步推进。',
            'system_upgrade': '系统已有功能待优化升级，结合未来工厂建设要求开展功能扩展。',
            'confirmed': '企业已明确建设计划的项目，需聚焦核心目标推进落地实施，确保与未来工厂建设要求对齐。',
        }

        for cat_key, cat_title in categories.items():
            self.doc.add_heading(cat_title, level=2)
            items = self.db.query(InvestmentItem).filter(
                InvestmentItem.project_id == self.project.id,
                InvestmentItem.category == cat_key
            ).order_by(InvestmentItem.sort_order).all()

            if not items:
                self.doc.add_paragraph('暂无投资项目。')
                continue

            # 类别说明
            self.doc.add_paragraph(category_intro.get(cat_key, ''))

            for i, item in enumerate(items, 1):
                budget_str = ''
                if item.budget_max and item.budget_min:
                    budget_str = f'（投入预算：{item.budget_min}-{item.budget_max}万元）'
                elif item.budget_min:
                    budget_str = f'（投入预算：{item.budget_min}万元）'

                p = self.doc.add_paragraph()
                p.add_run(f'{i}、{item.title}{budget_str}').bold = True

                if item.description:
                    self.doc.add_paragraph(item.description)

        # ---- 投资项目汇总表 ----
        all_items = self.db.query(InvestmentItem).filter(
            InvestmentItem.project_id == self.project.id
        ).order_by(InvestmentItem.sort_order).all()

        if all_items:
            self.doc.add_paragraph('')
            self.doc.add_heading('投资项目汇总表', level=2)
            table = self.doc.add_table(rows=1, cols=5, style='Light Grid Accent 1')
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            headers = ['序号', '类别', '项目名称', '说明', '预算（万元）']
            for i, h in enumerate(headers):
                self._set_cell_font(table.rows[0].cells[i], h, bold=True, size=9)

            idx = 1
            for item in all_items:
                row = table.add_row()
                self._set_cell_font(row.cells[0], str(idx), size=9)
                self._set_cell_font(row.cells[1], item.category, size=9)
                self._set_cell_font(row.cells[2], item.title, size=9)
                self._set_cell_font(row.cells[3], item.description or '', size=9)
                budget = ''
                if item.budget_max and item.budget_min:
                    budget = f'{item.budget_min}-{item.budget_max}'
                elif item.budget_min:
                    budget = str(item.budget_min)
                self._set_cell_font(row.cells[4], budget, size=9)
                idx += 1
